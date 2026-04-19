"""Generate the Rey Capital AI Bot Setup & Operations Guide as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Custom styles ──
style = doc.styles

# Heading 1 — dark navy blue
h1 = style['Heading 1']
h1.font.color.rgb = RGBColor(0, 0x52, 0xCC)
h1.font.size = Pt(22)
h1.font.bold = True

# Heading 2
h2 = style['Heading 2']
h2.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)
h2.font.size = Pt(16)
h2.font.bold = True

# Heading 3
h3 = style['Heading 3']
h3.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
h3.font.size = Pt(13)
h3.font.bold = True

# Normal
normal = style['Normal']
normal.font.size = Pt(11)
normal.font.name = 'Calibri'

# ── Helper functions ──
def add_code_block(text):
    """Add a styled code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    # Add shading via XML
    from docx.oxml.ns import qn
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F4F8'
    })
    shading.append(shd)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_warning(text):
    p = doc.add_paragraph()
    run = p.add_run("[!] " + text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run("[NOTE] " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ═══════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════

# Add logo if available
logo_path = os.path.join(os.path.dirname(__file__), 'static', 'ReyCapital_Logo.png')
if os.path.exists(logo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(2.5))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('Rey Capital AI Bot')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 0x52, 0xCC)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Setup & Operations Guide')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run(f'Version 1.0 | {datetime.now().strftime("%B %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
run = p.add_run('Powered by Gemma 4 AI + MetaTrader 5')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. System Overview',
    '2. Architecture',
    '3. Fresh Windows Server Setup (Step-by-Step)',
    '4. Configuration Reference',
    '5. Running the Bot',
    '6. Dashboard Guide',
    '7. Trading Symbols & Parameters',
    '8. Self-Learning Adaptive System',
    '9. Monitoring & Logs',
    '10. Troubleshooting',
    '11. AWS EC2 Deployment',
    '12. File Structure Reference',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  1. SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════════════

doc.add_heading('1. System Overview', level=1)

doc.add_paragraph(
    'Rey Capital AI Bot is an automated cryptocurrency scalping system that uses '
    'Google\'s Gemma 4 large language model (running locally via Ollama) to analyze '
    '30+ technical indicators on 1-minute candle data and execute profitable trades '
    'on MetaTrader 5.'
)

doc.add_heading('Key Features', level=2)
features = [
    ('AI-Powered Decisions: ', 'Gemma 4 analyzes 30+ indicators including Ichimoku Cloud, Supertrend, MACD, RSI, Bollinger Bands, and candlestick patterns'),
    ('Auto-Execution: ', 'Places BUY/SELL orders directly on MetaTrader 5 with calculated SL/TP'),
    ('Self-Learning: ', 'Tracks trade outcomes and builds adaptive context to improve future decisions'),
    ('Real-Time Dashboard: ', 'Professional web UI with dark/light theme, WebSocket updates, and trade monitoring'),
    ('Multi-Symbol: ', 'Simultaneously analyzes BTCUSD, ETHUSD, LTCUSD, XRPUSD, SOLUSD'),
    ('Risk Management: ', 'Per-symbol cooldowns, dynamic confidence thresholds, ATR-based position sizing, daily loss limits'),
]
for bold, text in features:
    add_bullet(text, bold)

doc.add_heading('Technology Stack', level=2)
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = table.rows[0].cells
headers[0].text = 'Component'
headers[1].text = 'Technology'
data = [
    ('AI Model', 'Gemma 4 (8B params, Q4_K_M quantized) via Ollama'),
    ('Broker', 'MetaTrader 5 (CFD crypto trading)'),
    ('Indicators', 'pandas_ta library (30+ technical indicators)'),
    ('Dashboard', 'Flask + Flask-SocketIO + WebSocket'),
    ('Language', 'Python 3.11+ (tested on 3.14)'),
    ('OS', 'Windows 10/11 or Windows Server 2019+'),
]
for i, (comp, tech) in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = comp
    row[1].text = tech

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  2. ARCHITECTURE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('2. Architecture', level=1)

doc.add_paragraph('The bot follows a simple pipeline that runs every 60 seconds:')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run(
    'MT5 (1M Candles)  -->  pandas_ta (30+ Indicators)  -->  Gemma 4 (AI Decision)\n'
    '       |                                                        |\n'
    '       v                                                        v\n'
    '  Price Data                                              BUY / SELL / HOLD\n'
    '                                                                |\n'
    '                                                                v\n'
    '                                                    MT5 Order Execution\n'
    '                                                                |\n'
    '                                                                v\n'
    '                                                    Trade Outcome Tracker\n'
    '                                                                |\n'
    '                                                                v\n'
    '                                                    Self-Learning Loop\n'
    '                                                   (Adaptive Context)'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('Data Flow Per Cycle', level=2)
steps = [
    'Fetch 500 candles from MT5 for each crypto symbol (BTCUSD, ETHUSD, LTCUSD, XRPUSD, SOLUSD)',
    'Calculate 30+ technical indicators using pandas_ta (EMAs, Ichimoku, MACD, RSI, Supertrend, etc.)',
    'Format all indicators into a structured prompt and send to Gemma 4 via Ollama',
    'Gemma 4 analyzes and returns JSON: {action, confidence, sl_distance_atr, tp_distance_atr, reason}',
    'Risk manager validates: confidence > threshold, daily loss limit, symbol cooldown, position limits',
    'If approved: calculate lot size using MT5 symbol info, compute SL/TP prices from ATR, place order',
    'Monitor open positions; when closed (SL/TP hit), record outcome for self-learning',
    'Periodically analyze trade history and update adaptive context (Gemma learns from mistakes)',
]
for i, step in enumerate(steps):
    add_numbered(step)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  3. FRESH WINDOWS SERVER SETUP
# ═══════════════════════════════════════════════════════════════

doc.add_heading('3. Fresh Windows Server Setup (Step-by-Step)', level=1)

doc.add_paragraph(
    'This guide covers setting up the bot from scratch on a clean Windows machine '
    '(Windows 10/11, Windows Server 2019/2022, or AWS EC2 Windows instance).'
)

# ── Step 1: Python ──
doc.add_heading('Step 1: Install Python 3.11+', level=2)
doc.add_paragraph('Download and install Python from https://www.python.org/downloads/')
add_bullet('Choose Python 3.11, 3.12, 3.13, or 3.14')
add_bullet('IMPORTANT: Check "Add Python to PATH" during installation')
add_bullet('IMPORTANT: Check "Install for all users" if on a server')
doc.add_paragraph('Verify installation:')
add_code_block('python --version')

# ── Step 2: MetaTrader 5 ──
doc.add_heading('Step 2: Install MetaTrader 5', level=2)
doc.add_paragraph('Download MT5 from your broker or https://www.metatrader5.com/en/download')
add_numbered('Install MetaTrader 5 (default location is fine)')
add_numbered('Launch MT5 and log into your broker account (demo or live)')
add_numbered('Go to Tools > Options > Expert Advisors')
add_numbered('Check "Allow automated trading"')
add_numbered('Click the "AutoTrading" button in the toolbar (must be GREEN)')
add_numbered('In Market Watch, right-click and "Show All" to ensure crypto symbols are visible')
add_warning('AutoTrading MUST be enabled (green button) or all trade orders will fail with error 10027.')

doc.add_paragraph('Verify your crypto symbols are available in Market Watch:')
add_bullet('BTCUSD, ETHUSD, LTCUSD, XRPUSD, SOLUSD')
add_note('Symbol names may vary by broker. Check your broker\'s symbol naming convention.')

# ── Step 3: Ollama + Gemma 4 ──
doc.add_heading('Step 3: Install Ollama & Gemma 4 Model', level=2)
doc.add_paragraph('Ollama runs AI models locally on your machine. No cloud API needed.')

add_numbered('Download Ollama from https://ollama.com/download (Windows installer)')
add_numbered('Install Ollama (it runs as a background service automatically)')
add_numbered('Open Command Prompt or PowerShell and pull Gemma 4:')
add_code_block('ollama pull gemma4')
add_numbered('Verify Ollama is running:')
add_code_block('curl http://localhost:11434/api/tags')

add_note('Gemma 4 (Q4_K_M) requires ~6GB RAM and ~9.6GB disk space. A GPU (NVIDIA with 8GB+ VRAM) is recommended for faster inference but not required.')
add_note('Ollama auto-starts on boot. If it\'s not running, start it with: ollama serve')

# ── Step 4: Bot Files ──
doc.add_heading('Step 4: Copy Bot Files', level=2)
doc.add_paragraph('Copy the entire gemma_trader folder to your server:')
add_code_block(r'D:\Projects\AlgoStrategies\AlgoStrategies\execution\gemma_trader')

doc.add_paragraph('Or clone from your Git repository:')
add_code_block('git clone <your-repo-url>\ncd AlgoStrategies/execution/gemma_trader')

# ── Step 5: Python Dependencies ──
doc.add_heading('Step 5: Install Python Dependencies', level=2)
add_code_block(
    'cd D:\\Projects\\AlgoStrategies\\AlgoStrategies\\execution\\gemma_trader\n'
    'pip install -r requirements.txt'
)

doc.add_paragraph('Required packages (from requirements.txt):')
table = doc.add_table(rows=10, cols=2)
table.style = 'Light Shading Accent 1'
headers = table.rows[0].cells
headers[0].text = 'Package'
headers[1].text = 'Purpose'
pkgs = [
    ('flask', 'Web dashboard'),
    ('flask-socketio', 'Real-time WebSocket updates'),
    ('pyyaml', 'Configuration file parsing'),
    ('requests', 'HTTP client for Ollama API'),
    ('pandas', 'Data manipulation'),
    ('pandas_ta', 'Technical indicators (30+)'),
    ('numpy', 'Numerical computations'),
    ('MetaTrader5', 'MT5 Python API for data & trading'),
    ('schedule', 'Task scheduling'),
]
for i, (pkg, purpose) in enumerate(pkgs):
    row = table.rows[i + 1].cells
    row[0].text = pkg
    row[1].text = purpose

add_note('If using Python 3.14, pandas_ta may need: pip install pandas_ta --no-deps followed by pip install tqdm scipy stockstats. A numba stub may also be needed.')

# ── Step 6: Configuration ──
doc.add_heading('Step 6: Configure the Bot', level=2)
doc.add_paragraph('Edit config.yaml with your settings:')
add_code_block(
    'notepad config.yaml'
)

doc.add_paragraph('Key settings to update:')
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Shading Accent 1'
h = table.rows[0].cells
h[0].text = 'Setting'
h[1].text = 'Default'
h[2].text = 'Description'
settings = [
    ('ollama.url', 'http://localhost:11434/api/generate', 'Ollama API endpoint'),
    ('ollama.model', 'gemma4', 'Model name (must match ollama pull)'),
    ('trading.mode', 'live', '"live" for real trades, "paper" for simulation'),
    ('trading.confidence_threshold', '0.65', 'Minimum AI confidence to execute trade'),
    ('trading.allowed_symbols', 'BTCUSD, ETHUSD, ...', 'Crypto symbols to trade'),
    ('broker.mt5.login', '0', 'Your MT5 account number'),
    ('broker.mt5.password', '(empty)', 'Your MT5 password (optional if already logged in)'),
]
for i, (s, d, desc) in enumerate(settings):
    row = table.rows[i + 1].cells
    row[0].text = s
    row[1].text = d
    row[2].text = desc

# ── Step 7: Create logs directory ──
doc.add_heading('Step 7: Create Logs Directory', level=2)
add_code_block('mkdir logs')
doc.add_paragraph('Initialize empty JSON log files:')
add_code_block(
    'echo [] > logs\\gemma_decisions.json\n'
    'echo [] > logs\\trades.json\n'
    'echo [] > logs\\trade_outcomes.json\n'
    'echo [] > logs\\parameter_adjustments.json'
)

# ── Step 8: Start ──
doc.add_heading('Step 8: Start the Bot', level=2)
add_code_block(
    'cd D:\\Projects\\AlgoStrategies\\AlgoStrategies\\execution\\gemma_trader\n'
    'python run.py'
)

doc.add_paragraph('You should see:')
add_code_block(
    '+===========================================================+\n'
    '|               REY CAPITAL AI BOT                           |\n'
    '+===========================================================+\n'
    '|  Dashboard:  http://localhost:8050                         |\n'
    '|  Mode:       LIVE                                         |\n'
    '|  Model:      gemma4                                       |\n'
    '|  Symbols:    BTCUSD, ETHUSD, LTCUSD, XRPUSD, SOLUSD     |\n'
    '|  Interval:   1m                                           |\n'
    '|  Trading:    ENABLED                                      |\n'
    '+===========================================================+'
)

doc.add_paragraph('Open your browser to http://localhost:8050 to see the dashboard.')

# ── Step 9: Verify ──
doc.add_heading('Step 9: Verify Everything Works', level=2)
doc.add_paragraph('Checklist:')
add_numbered('MT5 shows "MT5 connected | Account: XXXXX | Balance: XXXXX" in console')
add_numbered('Bot starts analyzing symbols: "Analyzing BTCUSD on 1m"')
add_numbered('Gemma responds: "Gemma: SELL (70%) - reason..."')
add_numbered('Trades execute: "[OK] MT5 SELL 8.0 BTCUSD @ 70976.0"')
add_numbered('Dashboard loads at http://localhost:8050 with Rey Capital branding')
add_numbered('Dashboard shows live decisions, trades, and statistics')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  4. CONFIGURATION REFERENCE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('4. Configuration Reference (config.yaml)', level=1)

doc.add_heading('Server', level=2)
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Shading Accent 1'
h = table.rows[0].cells
h[0].text = 'Key'; h[1].text = 'Default'; h[2].text = 'Description'
for i, (k, v, d) in enumerate([
    ('server.host', '0.0.0.0', 'Dashboard bind address'),
    ('server.port', '8050', 'Dashboard port'),
]):
    r = table.rows[i+1].cells; r[0].text = k; r[1].text = v; r[2].text = d

doc.add_heading('Ollama (AI Model)', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Shading Accent 1'
h = table.rows[0].cells
h[0].text = 'Key'; h[1].text = 'Default'; h[2].text = 'Description'
for i, (k, v, d) in enumerate([
    ('ollama.url', 'http://localhost:11434/api/generate', 'Ollama endpoint'),
    ('ollama.model', 'gemma4', 'Model name'),
    ('ollama.temperature', '0.1', 'Creativity (lower = more consistent)'),
    ('ollama.timeout', '120', 'Request timeout in seconds'),
]):
    r = table.rows[i+1].cells; r[0].text = k; r[1].text = v; r[2].text = d

doc.add_heading('Trading', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Shading Accent 1'
h = table.rows[0].cells
h[0].text = 'Key'; h[1].text = 'Default'; h[2].text = 'Description'
for i, (k, v, d) in enumerate([
    ('trading.mode', 'live', '"live" = real trades, "paper" = simulated'),
    ('trading.confidence_threshold', '0.65', 'Min confidence to trade (0.0-1.0)'),
    ('trading.max_position_size_pct', '1.0', 'Max % of balance per trade'),
    ('trading.max_open_trades', '5', 'Max simultaneous open positions'),
    ('trading.cooldown_minutes', '3', 'Min minutes between trades on same symbol'),
]):
    r = table.rows[i+1].cells; r[0].text = k; r[1].text = v; r[2].text = d

doc.add_heading('Risk Management', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Shading Accent 1'
h = table.rows[0].cells
h[0].text = 'Key'; h[1].text = 'Default'; h[2].text = 'Description'
for i, (k, v, d) in enumerate([
    ('risk_management.stop_loss_atr_multiplier', '1.0', 'SL distance = ATR x this'),
    ('risk_management.take_profit_atr_multiplier', '1.5', 'TP distance = ATR x this'),
    ('risk_management.max_daily_loss_pct', '5.0', 'Stop trading if daily loss > 5%'),
    ('risk_management.max_drawdown_pct', '10.0', 'Stop trading if drawdown > 10%'),
]):
    r = table.rows[i+1].cells; r[0].text = k; r[1].text = v; r[2].text = d

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  5. RUNNING THE BOT
# ═══════════════════════════════════════════════════════════════

doc.add_heading('5. Running the Bot', level=1)

doc.add_heading('Quick Start', level=2)
add_code_block(
    'cd D:\\Projects\\AlgoStrategies\\AlgoStrategies\\execution\\gemma_trader\n'
    'python run.py'
)

doc.add_heading('Command-Line Options', level=2)
table = doc.add_table(rows=7, cols=3)
table.style = 'Light Shading Accent 1'
h = table.rows[0].cells
h[0].text = 'Flag'; h[1].text = 'Example'; h[2].text = 'Description'
for i, (f, e, d) in enumerate([
    ('--mode', 'python run.py --mode paper', 'Paper trading (no real orders)'),
    ('--port', 'python run.py --port 9000', 'Custom dashboard port'),
    ('--symbols', 'python run.py --symbols BTCUSD ETHUSD', 'Trade only specific symbols'),
    ('--interval', 'python run.py --interval 5m', 'Change candle timeframe'),
    ('--no-trade', 'python run.py --no-trade', 'Dashboard only, no trading'),
    ('--config', 'python run.py --config my_config.yaml', 'Use alternate config file'),
]):
    r = table.rows[i+1].cells; r[0].text = f; r[1].text = e; r[2].text = d

doc.add_heading('Prerequisites (Must Be Running)', level=2)
add_numbered('MetaTrader 5 desktop - logged in, AutoTrading ON (green button)')
add_numbered('Ollama service - running with Gemma 4 model loaded')

doc.add_heading('Run as Background Service (Windows)', level=2)
doc.add_paragraph('Option A: Batch file')
add_code_block(
    '@echo off\n'
    'cd /d D:\\Projects\\AlgoStrategies\\AlgoStrategies\\execution\\gemma_trader\n'
    'python run.py >> logs\\bot_output.log 2>&1'
)

doc.add_paragraph('Option B: Windows Task Scheduler (auto-start on boot)')
add_numbered('Open Task Scheduler (taskschd.msc)')
add_numbered('Click "Create Basic Task"')
add_numbered('Name: "Rey Capital AI Bot"')
add_numbered('Trigger: "When the computer starts"')
add_numbered('Action: Start a program')
add_numbered('Program: python.exe')
add_numbered('Arguments: D:\\Projects\\...\\execution\\gemma_trader\\run.py')
add_numbered('Start in: D:\\Projects\\...\\execution\\gemma_trader')

doc.add_paragraph('Option C: NSSM (Non-Sucking Service Manager) for Windows Services')
add_code_block(
    'nssm install ReyCaptialBot python.exe run.py\n'
    'nssm set ReyCaptialBot AppDirectory D:\\Projects\\...\\execution\\gemma_trader\n'
    'nssm start ReyCaptialBot'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  6. DASHBOARD GUIDE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('6. Dashboard Guide', level=1)
doc.add_paragraph('Access the dashboard at http://localhost:8050 (or your custom port).')

doc.add_heading('Dashboard Sections', level=2)
sections = [
    ('Stats Row: ', 'Shows account balance, total trades, win rate, daily P&L, and active positions'),
    ('Symbol Cards: ', 'Current status of each crypto symbol with latest AI decision and confidence'),
    ('AI Decisions Timeline: ', 'Real-time feed of Gemma\'s analysis for each symbol'),
    ('Trade Log: ', 'History of all executed trades with entry/exit prices, SL/TP, and P&L'),
    ('AI Learning Card: ', 'Shows what the AI has learned from past trades (adaptive context)'),
]
for bold, text in sections:
    add_bullet(text, bold)

doc.add_heading('Theme Toggle', level=2)
doc.add_paragraph('Click the sun/moon icon in the header to switch between dark and light themes. Your preference is saved in the browser.')

doc.add_heading('API Endpoints', level=2)
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Shading Accent 1'
h = table.rows[0].cells
h[0].text = 'Endpoint'; h[1].text = 'Description'
for i, (ep, desc) in enumerate([
    ('GET /', 'Dashboard HTML page'),
    ('GET /api/health', 'System status, balance, Ollama status'),
    ('GET /api/decisions', 'Recent AI decisions (last 100)'),
    ('GET /api/trades', 'Recent executed trades'),
    ('GET /api/learning', 'Adaptive context, win rate, outcomes'),
    ('GET /api/pnl', 'Real-time P&L from MT5 open positions'),
]):
    r = table.rows[i+1].cells; r[0].text = ep; r[1].text = desc

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  7. TRADING SYMBOLS & PARAMETERS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('7. Trading Symbols & Parameters', level=1)

doc.add_heading('Supported Crypto Symbols', level=2)
table = doc.add_table(rows=6, cols=5)
table.style = 'Light Shading Accent 1'
h = table.rows[0].cells
h[0].text = 'Symbol'; h[1].text = 'Name'; h[2].text = 'Min Lot'; h[3].text = 'Contract Size'; h[4].text = 'Notes'
symbols = [
    ('BTCUSD', 'Bitcoin', '0.001', '1.0', 'Highest liquidity, drives market'),
    ('ETHUSD', 'Ethereum', '0.001', '10.0', 'DeFi catalyst, correlates with BTC'),
    ('LTCUSD', 'Litecoin', '0.001', '100.0', 'Fast mover, leads BTC moves'),
    ('XRPUSD', 'Ripple', '0.001', '5000.0', 'News-driven, regulatory sensitive'),
    ('SOLUSD', 'Solana', '0.01', '100.0', 'High-beta, amplifies BTC 2-3x'),
]
for i, (s, n, ml, cs, note) in enumerate(symbols):
    r = table.rows[i+1].cells
    r[0].text = s; r[1].text = n; r[2].text = ml; r[3].text = cs; r[4].text = note

add_note('Lot sizes and contract sizes are from CFI3-Demo broker. Your broker may differ.')

doc.add_heading('Indicators Used (30+)', level=2)
doc.add_heading('Trend Indicators', level=3)
for ind in ['EMA (9, 20, 50, 200)', 'SMA (20, 50)', 'Ichimoku Cloud (Tenkan, Kijun, Span A/B, Cloud Color)', 'Supertrend (10, 3)', 'Parabolic SAR', 'ADX + DI+/DI- (14)', 'VWAP']:
    add_bullet(ind)

doc.add_heading('Momentum Indicators', level=3)
for ind in ['RSI (14)', 'MACD (12, 26, 9) + Histogram', 'Stochastic RSI (14)', 'Stochastic Oscillator (14, 3)', 'CCI (20)', 'Williams %R (14)', 'ROC (10)', 'MFI (14)']:
    add_bullet(ind)

doc.add_heading('Volatility Indicators', level=3)
for ind in ['ATR (14)', 'Bollinger Bands (20, 2.0) - Upper, Mid, Lower, Width, %B']:
    add_bullet(ind)

doc.add_heading('Volume Indicators', level=3)
for ind in ['Volume with SMA(20) ratio', 'Volume trend classification (HIGH/LOW/ABOVE_AVG/BELOW_AVG)', 'VWAP']:
    add_bullet(ind)

doc.add_heading('Price Action', level=3)
for ind in ['Last 5 candles summary', 'Candlestick patterns (Doji, Hammer, Engulfing, etc.)', 'Nearest support/resistance levels']:
    add_bullet(ind)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  8. SELF-LEARNING ADAPTIVE SYSTEM
# ═══════════════════════════════════════════════════════════════

doc.add_heading('8. Self-Learning Adaptive System', level=1)

doc.add_paragraph(
    'The bot continuously learns from its own trading performance and adjusts its behavior. '
    'This is the key innovation that makes Rey Capital AI Bot improve over time.'
)

doc.add_heading('How It Works', level=2)

doc.add_heading('8.1 Trade Outcome Tracking', level=3)
doc.add_paragraph(
    'When a trade hits SL (loss) or TP (profit), the full context is recorded: '
    'entry indicators, Gemma\'s reasoning, confidence, outcome, P&L, and duration. '
    'Stored in logs/trade_outcomes.json.'
)

doc.add_heading('8.2 Performance Analysis', level=3)
doc.add_paragraph(
    'After every 10 trades, the TradeReviewer analyzes: win rate, avg win/loss, '
    'profit factor, per-symbol performance, and indicator pattern correlations. '
    'It identifies which indicator combinations lead to wins vs losses.'
)

doc.add_heading('8.3 Adaptive Prompt Engineering', level=3)
doc.add_paragraph(
    'Based on the analysis, an adaptive context is generated and prepended to '
    'Gemma\'s system prompt. Example lessons:'
)
add_code_block(
    'LESSONS FROM RECENT TRADES (auto-updated):\n'
    '- Avoid BUY on SOLUSD when RSI > 60 and volume is BELOW_AVG (lost 4/5)\n'
    '- BTCUSD SELL signals with ADX < 20 are unreliable (3 losses)\n'
    '- ETHUSD performs best when MACD histogram increasing AND volume HIGH\n'
    '- Current win rate: 58% - maintain discipline\n'
    '- Best: BTCUSD BUY signals (67% win rate)\n'
    '- Worst: XRPUSD SELL signals (25% win rate)'
)

doc.add_heading('8.4 Dynamic Parameter Adjustment', level=3)
add_bullet('Win rate < 40% over last 20 trades: confidence threshold auto-raises by 0.05')
add_bullet('Win rate > 60%: threshold can lower by 0.02 (minimum 0.50)')
add_bullet('3+ consecutive losses on a symbol: 15-minute cooldown for that symbol')
add_bullet('All adjustments logged in logs/parameter_adjustments.json')

doc.add_heading('8.5 Weekly Meta-Review', level=3)
doc.add_paragraph(
    'Periodically, Gemma reviews its own trade history and generates refined trading rules. '
    'This is "Gemma reviewing Gemma" - a meta-learning loop that produces increasingly '
    'refined decision-making over time.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  9. MONITORING & LOGS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('9. Monitoring & Logs', level=1)

doc.add_heading('Log Files', level=2)
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Shading Accent 1'
h = table.rows[0].cells
h[0].text = 'File'; h[1].text = 'Description'
for i, (f, d) in enumerate([
    ('logs/gemma_decisions.json', 'All AI decisions with timestamps, symbols, indicators'),
    ('logs/trades.json', 'Executed trades with entry price, SL, TP, qty'),
    ('logs/trade_outcomes.json', 'Closed trades with P&L, win/loss, indicator snapshots'),
    ('logs/adaptive_context.txt', 'Current lessons learned by the AI'),
    ('logs/parameter_adjustments.json', 'History of threshold/parameter auto-adjustments'),
]):
    r = table.rows[i+1].cells; r[0].text = f; r[1].text = d

doc.add_heading('Console Output', level=2)
doc.add_paragraph('The bot logs all activity to the console with timestamps:')
add_bullet('Cycle start with balance and open positions')
add_bullet('Per-symbol analysis: RSI, MACD, Trend, Ichimoku signal, Volume')
add_bullet('Gemma decisions with confidence and reasoning')
add_bullet('Trade execution results (success or error)')
add_bullet('Position close events with P&L')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  10. TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════

doc.add_heading('10. Troubleshooting', level=1)

problems = [
    (
        'Error 10027: AutoTrading disabled by client',
        'Enable AutoTrading in MT5:\n'
        '1. Click the "AutoTrading" button in MT5 toolbar (turn it GREEN)\n'
        '2. Go to Tools > Options > Expert Advisors > Check "Allow automated trading"'
    ),
    (
        'Error 10030: Unsupported filling mode',
        'The bot auto-detects the correct filling mode (FOK/IOC/RETURN) from symbol info. '
        'If this error persists, check that the symbol is tradeable on your broker.'
    ),
    (
        'Error 10016: Invalid stops',
        'SL/TP prices are too close to the current price (within spread + freeze level). '
        'The bot calculates minimum distances from spread and freeze level. '
        'If ATR is very small, the bot ensures SL/TP are at least spread + freeze + 10 points away.'
    ),
    (
        'Gemma returns empty response',
        'Increase num_predict in config.yaml to 8192. Gemma 4 uses ~500 tokens for internal '
        'thinking before producing visible output. With num_predict < 500, response may be empty.'
    ),
    (
        'Cannot connect to Ollama',
        'Ensure Ollama is running: "ollama serve" or check Windows Services for "Ollama".\n'
        'Verify: curl http://localhost:11434/api/tags\n'
        'If running on a different machine, update ollama.url in config.yaml.'
    ),
    (
        'MT5 not connected / Cannot get candles',
        '1. Ensure MetaTrader 5 desktop is open and logged in\n'
        '2. Check that symbols are in Market Watch (right-click > Show All)\n'
        '3. Ensure the MetaTrader5 Python package is installed: pip install MetaTrader5'
    ),
    (
        'UnicodeDecodeError on config.yaml',
        'Ensure config.yaml uses ASCII-only characters (no Unicode box-drawing characters). '
        'Or open with encoding="utf-8" (already handled in the code).'
    ),
    (
        'pandas_ta / numba import error on Python 3.14',
        'Install pandas_ta without dependencies: pip install pandas_ta --no-deps\n'
        'Then install: pip install tqdm scipy stockstats\n'
        'Create a numba stub if needed (see deployment scripts).'
    ),
    (
        'Dashboard shows "Error reading gemma_decisions.json"',
        'The JSON file may have a BOM (byte order mark). Delete and recreate:\n'
        'echo [] > logs\\gemma_decisions.json'
    ),
]

for problem, solution in problems:
    doc.add_heading(problem, level=3)
    doc.add_paragraph(solution)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  11. AWS EC2 DEPLOYMENT
# ═══════════════════════════════════════════════════════════════

doc.add_heading('11. AWS EC2 Deployment (Windows Server)', level=1)

doc.add_heading('Recommended Instance', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Shading Accent 1'
h = table.rows[0].cells
h[0].text = 'Setting'; h[1].text = 'Recommendation'
for i, (s, r) in enumerate([
    ('Instance Type', 'g4dn.xlarge (4 vCPU, 16 GB RAM, T4 GPU) or t3.xlarge (no GPU)'),
    ('AMI', 'Windows Server 2022 Base'),
    ('Storage', '100 GB SSD (GP3)'),
    ('Security Group', 'Inbound: RDP (3389), Dashboard (8050), Ollama (11434)'),
]):
    r2 = table.rows[i+1].cells; r2[0].text = s; r2[1].text = r

doc.add_heading('Setup Steps', level=2)
add_numbered('Launch Windows Server EC2 instance')
add_numbered('RDP into the instance')
add_numbered('Install Python 3.11+ (from python.org, check "Add to PATH")')
add_numbered('Install MetaTrader 5 and log into your broker account')
add_numbered('Install Ollama from ollama.com/download')
add_numbered('Pull Gemma 4: ollama pull gemma4')
add_numbered('Copy bot files to the server (git clone or SCP)')
add_numbered('Install dependencies: pip install -r requirements.txt')
add_numbered('Configure config.yaml with your MT5 credentials')
add_numbered('Enable AutoTrading in MT5')
add_numbered('Start the bot: python run.py')
add_numbered('Use NSSM to install as Windows Service for persistence')

add_warning('MT5 requires a desktop session (GUI). On EC2, keep the RDP session alive or use a virtual display solution. Consider setting up "Keep RDP alive" via Group Policy.')

doc.add_heading('Firewall Rules (PowerShell)', level=2)
add_code_block(
    '# Allow dashboard access\n'
    'New-NetFirewallRule -DisplayName "Rey Capital Dashboard" -Direction Inbound -Port 8050 -Protocol TCP -Action Allow\n\n'
    '# Allow Ollama (if accessing remotely)\n'
    'New-NetFirewallRule -DisplayName "Ollama API" -Direction Inbound -Port 11434 -Protocol TCP -Action Allow'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  12. FILE STRUCTURE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('12. File Structure Reference', level=1)

add_code_block(
    'gemma_trader/\n'
    '|-- run.py                    # Main entry point (start here)\n'
    '|-- config.yaml               # All configuration settings\n'
    '|-- local_trader.py           # Trading engine + indicator calculator\n'
    '|-- gemma_analyzer.py         # Gemma 4 AI integration + prompts\n'
    '|-- risk_manager.py           # Position sizing + risk controls\n'
    '|-- broker_bridge.py          # MT5/Binance/Paper trade execution\n'
    '|-- mt5_data_feed.py          # MT5 candle data fetching\n'
    '|-- trade_reviewer.py         # Self-learning + performance analysis\n'
    '|-- dashboard.py              # Flask web dashboard + API\n'
    '|-- requirements.txt          # Python dependencies\n'
    '|\n'
    '|-- templates/\n'
    '|   |-- dashboard.html        # Dashboard UI (dark/light theme)\n'
    '|\n'
    '|-- static/\n'
    '|   |-- ReyCapital_Logo.png       # Logo (light theme)\n'
    '|   |-- ReyCapital_Logo_White.png # Logo (dark theme)\n'
    '|   |-- ReyCapital_Icon.png       # Favicon\n'
    '|\n'
    '|-- logs/\n'
    '|   |-- gemma_decisions.json      # AI decision history\n'
    '|   |-- trades.json               # Executed trade log\n'
    '|   |-- trade_outcomes.json       # Trade outcomes for learning\n'
    '|   |-- adaptive_context.txt      # Current AI lessons\n'
    '|   |-- parameter_adjustments.json # Auto-adjustment history\n'
    '|\n'
    '|-- deploy/\n'
    '    |-- setup_windows_ec2.ps1     # AWS EC2 setup script\n'
    '    |-- setup_local.bat           # Local Windows setup script'
)

doc.add_paragraph('')

# ── Footer ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('---')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Rey Capital AI Bot | Powered by Gemma 4')
run.font.color.rgb = RGBColor(0, 0x52, 0xCC)
run.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Document generated on {datetime.now().strftime("%B %d, %Y")}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'Rey_Capital_AI_Bot_Setup_Guide.docx')
doc.save(output_path)
print(f"Document saved to: {output_path}")
